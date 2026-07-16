#!/usr/bin/env python3
"""Render GUION_PRESENTACION.md + GUION_VIDEO.md into a single, visual Word doc.

Re-run whenever the guiones change to keep the .docx in sync.
Minimal Markdown support: #/##/### headings, | tables |, - bullets, > quotes,
**bold** and `code` inline.
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = "/home/sonda/Documentos/voctomix/presentacion_tfg"
OUT = os.path.join(HERE, "GUION_Presentacion_Voctomix.docx")
BLUE = RGBColor(0x08, 0x44, 0x7C)
CYAN = RGBColor(0x28, 0x99, 0xF9)
GRAY = RGBColor(0x46, 0x47, 0x47)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.font.color.rgb = GRAY


def add_runs(par, text):
    """Parse **bold** and `code` inline."""
    for tok in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True; r.font.color.rgb = BLUE
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"
        else:
            par.add_run(tok)


def render_md(path, doc):
    lines = open(path, encoding="utf-8").read().splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        # table
        if ln.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in block]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(t.rows[ri].cells):
                            cpar = t.rows[ri].cells[ci].paragraphs[0]
                            add_runs(cpar, cell)
                            for run in cpar.runs:
                                run.font.size = Pt(9)
                                if ri == 0:
                                    run.bold = True
            continue
        if ln.startswith("### "):
            h = doc.add_heading(ln[4:], level=3); h.runs[0].font.color.rgb = CYAN
        elif ln.startswith("## "):
            h = doc.add_heading(ln[3:], level=2); h.runs[0].font.color.rgb = BLUE
        elif ln.startswith("# "):
            h = doc.add_heading(ln[2:], level=1); h.runs[0].font.color.rgb = BLUE
        elif ln.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, ln.lstrip()[2:])
        elif ln.lstrip().startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, ln.lstrip()[2:])
            for r in p.runs:
                r.italic = True; r.font.color.rgb = CYAN
        else:
            p = doc.add_paragraph(); add_runs(p, ln)
        i += 1


title = doc.add_heading("Guion de la defensa — Voctomix 2.0", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Producción remota de vídeo en tiempo real con software libre  ·  "
                  "Martín Herranz Sánchez").alignment = WD_ALIGN_PARAGRAPH.CENTER

render_md(os.path.join(HERE, "GUION_PRESENTACION.md"), doc)
doc.add_page_break()
render_md(os.path.join(HERE, "GUION_VIDEO.md"), doc)

doc.save(OUT)
print("Saved:", OUT)
